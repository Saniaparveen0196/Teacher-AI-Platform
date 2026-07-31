# app/parsers/docx_parser.py
"""
Stage 1 — Document Intelligence: DOCX parser.
Word has real heading styles, so no regex guessing needed here.
"""
from docx import Document


def parse_docx(file_path: str) -> dict:
    doc = Document(file_path)
    sections = []
    tables = []
    current_section = {"heading": "Introduction", "level": 1, "text": ""}
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue  # skip blank spacing paragraphs

        full_text_parts.append(text)

        style_name = (para.style.name or "").lower() if para.style is not None else ""

        if style_name.startswith("heading") or style_name == "title":
            level = 1
            for ch in style_name:
                if ch.isdigit():
                    level = int(ch)
                    break
            if current_section["text"].strip():
                sections.append(current_section)
            current_section = {"heading": text, "level": level, "text": ""}
        else:
            current_section["text"] += text + "\n"

    if current_section["text"].strip():
        sections.append(current_section)

    for tbl in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        tables.append({"page": None, "rows": rows})

    raw_text = "\n".join(full_text_parts)
    return {
        "raw_text": raw_text,
        "sections": sections if sections else [{"heading": "Full Document", "level": 1, "text": raw_text}],
        "tables": tables,
        "images_or_figures": (
            [{"page": None, "caption": None, "count": len(doc.inline_shapes)}]
            if doc.inline_shapes else []
        ),
        "metadata": {
            "page_count": None,
            "word_count": len(raw_text.split()),
            "source_format": "docx",
        },
    }